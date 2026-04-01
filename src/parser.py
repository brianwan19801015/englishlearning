import re
from pathlib import Path
from docx import Document


def parse_docx(file_path: str) -> list[str]:
    """从Word文档提取句子"""
    doc = Document(file_path)
    sentences = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            sentences.append(text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    sentences.append(text)
    
    return sentences


def extract_exercises(sentences: list[str]) -> list[dict]:
    """从句子中提取填空题"""
    exercises = []
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        
        sentence = re.sub(r"^\d+\s+", "", sentence)
        sentence = re.sub(r"^\*\d+\.\s*", "", sentence)
        
        blank_match = re.search(r"_+", sentence)
        root_match = re.search(r"\(([^)]+)\)\s*$", sentence)
        
        if blank_match and root_match:
            root = root_match.group(1)
            context = sentence
            
            context = re.sub(r"\(([^)]+)\)\s*$", "", context)
            context = context.strip()
            context = context + " (" + root + ")"
            
            exercises.append({
                "context": context,
                "root": root
            })
    
    return exercises


def load_all_exercises(source_dir: str) -> list[dict]:
    """加载所有文档中的练习题"""
    all_exercises = []
    source_path = Path(source_dir)
    
    for docx_file in source_path.glob("*.docx"):
        if "~$" in docx_file.name:
            continue
        
        sentences = parse_docx(str(docx_file))
        exercises = extract_exercises(sentences)
        
        for ex in exercises:
            ex["source"] = docx_file.name
        
        all_exercises.extend(exercises)
    
    return all_exercises
